import os
import io
import math
import requests

import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from openai import OpenAI

# ───── SETUP ─────
OLLAMA_BASE_URL = "http://localhost:11434"
EMBED_MODEL     = "nomic-embed-text"
OPENAI_MODEL    = "gpt-4o-mini"

# Read from environment, default to empty string
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "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")

# ───── OLLAMA CLIENT ─────
class OllamaClient:
    def __init__(self, base_url=OLLAMA_BASE_URL):
        self.base_url = base_url

    def check_connection(self) -> bool:
        try:
            return requests.get(f"{self.base_url}/api/tags", timeout=5).status_code == 200
        except:
            return False

    def get_embeddings(self, texts: list[str], model=EMBED_MODEL) -> np.ndarray:
        """Batch-request embeddings for a list of texts.
        If the batched call returns too few results, fall back to per-text requests."""
        batch = []
        try:
            resp = requests.post(
                f"{self.base_url}/api/embeddings",
                json={"model": model, "prompt": texts},
                timeout=60
            )
            batch = resp.json().get("embeddings", [])
        except Exception:
            batch = []

        if len(batch) < len(texts):
            batch = []
            for txt in texts:
                single = []
                try:
                    r = requests.post(
                        f"{self.base_url}/api/embeddings",
                        json={"model": model, "prompt": txt},
                        timeout=30
                    )
                    single = r.json().get("embedding", [])
                except Exception:
                    single = []
                if not single:
                    single = [0.0] * 768
                batch.append(single)

        return np.array(batch)

ollama_client = OllamaClient()

# ───── UTILITIES ─────
def extract_text_from_docx(file) -> str:
    doc = Document(file)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = [
        c.text.strip()
        for t in doc.tables
        for r in t.rows
        for c in r.cells
        if c.text.strip()
    ]
    return "\n".join(paras + tables)

def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    na = a / np.linalg.norm(a)
    nb = b / np.linalg.norm(b)
    return float(np.dot(na, nb))

def generate_feedback(openai_client: OpenAI, jd: str, resume: str) -> str:
    prompt = f"""
You are an expert career coach. Given the following Job Description (JD) and Candidate Resume text,
provide concise feedback in two sections: "Strengths" and "Areas for Improvement",
focused on how well the resume matches the JD.

### Job Description:
{jd}

### Resume:
{resume}

Feedback:
- Strengths:
- Areas for Improvement:
"""
    resp = openai_client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    return resp.choices[0].message.content.strip()

# ───── STREAMLIT UI ─────
st.set_page_config("Resume-JD Matcher", layout="wide")
st.title("Resume vs JD Matcher (High‑Perf & Hybrid AI)")

# Sidebar for API Key
if not OPENAI_API_KEY:
    st.sidebar.markdown("## 🔑 OpenAI API Key")
    OPENAI_API_KEY = st.sidebar.text_input(
        "Paste your OpenAI key here",
        type="password",
        placeholder="sk-…",
        help="We only store this for the current session."
    ).strip()

openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# Pre‑flight Checks
if not ollama_client.check_connection():
    st.error("❌ Ollama server not running. Please start Ollama (`ollama serve`).")
    st.stop()

if not openai_client:
    st.error("❌ OpenAI API key is missing. Please set `OPENAI_API_KEY` in your environment or sidebar.")
    st.stop()

st.success("✅ Ollama & OpenAI clients ready")

# Initialize session state
if "results" not in st.session_state:
    st.session_state.results = []
if "jd_text" not in st.session_state:
    st.session_state.jd_text = ""
if "feedback" not in st.session_state:
    st.session_state.feedback = {}  # map resume_name -> feedback text

# File Upload
col1, col2 = st.columns(2)
with col1:
    jd_file = st.file_uploader("Upload Job Description (.docx)", type=["docx"])
with col2:
    resume_files = st.file_uploader(
        "Upload One or More Resumes (.docx)",
        type=["docx"],
        accept_multiple_files=True
    )

# Matching Trigger
if st.button("🔍 Match Resumes", disabled=not (jd_file and resume_files)):
    # Extract and store JD text
    st.session_state.jd_text = extract_text_from_docx(jd_file)

    # Extract resumes
    resumes_txt = [extract_text_from_docx(f) for f in resume_files]

    # Embed all texts in one batch
    all_texts = [st.session_state.jd_text] + resumes_txt
    all_embs = ollama_client.get_embeddings(all_texts)

    if all_embs.shape[0] < 1:
        st.error("Failed to retrieve embeddings for the Job Description. Aborting.")
        st.stop()

    jd_emb = all_embs[0]
    if all_embs.shape[0] < len(resumes_txt) + 1:
        st.warning("Warning: some resume embeddings failed; results may be incomplete.")
    resume_embs = all_embs[1:]

    # Compute similarity and collect results
    results = []
    for name, txt, emb in zip(
        [f.name for f in resume_files],
        resumes_txt,
        resume_embs
    ):
        sim_pct = round(cosine_sim(jd_emb, emb) * 100, 2)
        results.append({"resume_name": name, "match_perc": sim_pct, "text": txt})
    st.session_state.results = sorted(results, key=lambda r: r["match_perc"], reverse=True)

# Display Overall Summary
if st.session_state.results:
    st.markdown("---")
    st.subheader("Overall Match Summary")
    summary_df = pd.DataFrame([
        {"Resume": r["resume_name"], "Match %": f'{r["match_perc"]}%'}
        for r in st.session_state.results
    ])
    st.dataframe(summary_df, use_container_width=True, hide_index=True)

    # Detailed Per‑Resume Tabs
    st.markdown("---")
    st.subheader("Detailed Analysis & Feedback")
    tabs = st.tabs([r["resume_name"] for r in st.session_state.results])

    for tab, res in zip(tabs, st.session_state.results):
        with tab:
            st.metric("Match Percentage", f"{res['match_perc']}%")

            # Generate or display feedback
            if st.button("📝 Generate Feedback", key=f"fb_{res['resume_name']}"):
                with st.spinner("Generating feedback via OpenAI..."):
                    fb = generate_feedback(openai_client, st.session_state.jd_text, res["text"])
                    st.session_state.feedback[res["resume_name"]] = fb

            # Show feedback if available
            if res["resume_name"] in st.session_state.feedback:
                st.subheader("Feedback")
                st.markdown(st.session_state.feedback[res["resume_name"]].replace("\n", "\n\n"))

            # Download enhanced resume
            binary = resume_files[
                [f.name for f in resume_files].index(res["resume_name"])
            ].getvalue()
            doc = Document(io.BytesIO(binary))

            # If we have feedback, append it
            if res["resume_name"] in st.session_state.feedback:
                doc.add_page_break()
                doc.add_heading("AI Feedback", level=2)
                for line in st.session_state.feedback[res["resume_name"]].splitlines():
                    if line.strip():
                        doc.add_paragraph(line.strip("-• "), style=None)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            st.download_button(
                label="📥 Download Enhanced Resume with Feedback",
                data=buf,
                file_name=f"enhanced_{res['resume_name']}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
